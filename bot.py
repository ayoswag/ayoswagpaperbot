from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, MessageHandler, ContextTypes, filters, CommandHandler, CallbackQueryHandler
from docxtpl import DocxTemplate
from datetime import datetime
import os
import pdfkit
from docx import Document

# Получаем токен из переменных окружения (безопасно для сервера)
TOKEN = os.getenv("BOT_TOKEN")
if not TOKEN:
    raise ValueError("Переменная окружения BOT_TOKEN не установлена!")

# Словарь с типами договоров
CONTRACT_TYPES = {
    "exclusive": {
        "name": "Exclusive Rights",
        "template": "contract_template.docx",
        "description": "Эксклюзивные права на бит"
    },
    "rent": {
        "name": "Unlimited Lease",
        "template": "rent_template.docx",
        "description": "Аренда бита (Unlimited Lease)"
    }
}

def docx_to_html(docx_path):
    """Конвертирует DOCX в HTML для последующей конвертации в PDF"""
    doc = Document(docx_path)
    html = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            html.append(f"<p>{paragraph.text}</p>")
    # Добавляем таблицы
    for table in doc.tables:
        html.append("<table border='1' cellpadding='5'>")
        for row in table.rows:
            html.append("<tr>")
            for cell in row.cells:
                html.append(f"<td>{cell.text}</td>")
            html.append("</tr>")
        html.append("</table>")
    return "".join(html)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Команда /start - показывает меню выбора договора"""
    keyboard = [
        [
            InlineKeyboardButton("📄 Exclusive Rights", callback_data="exclusive"),
            InlineKeyboardButton("📄 Unlimited Lease", callback_data="rent")
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "👋 Привет! Я бот для создания договоров на биты.\n\n"
        "Выбери тип договора:",
        reply_markup=reply_markup
    )

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработка нажатия кнопок"""
    query = update.callback_query
    await query.answer()
    
    contract_type = query.data
    context.user_data['contract_type'] = contract_type
    
    template_name = CONTRACT_TYPES[contract_type]["name"]
    await query.edit_message_text(
        f"✅ Выбран: {template_name}\n\n"
        "Теперь отправь данные в таком порядке (каждое с новой строки):\n\n"
        "1️⃣ Название бита\n"
        "2️⃣ Имя артиста\n"
        "3️⃣ Никнейм\n"
        "4️⃣ Адрес\n"
        "5️⃣ Цена (только цифры)\n"
        "6️⃣ Процент паблишинга (только цифры)\n"
        "7️⃣ Дата (в формате MM.DD.YY)\n\n"
        "📝 Пример:\n"
        "Follow Me\n"
        "Lil Pump\n"
        "@lilpump\n"
        "Miami, USA\n"
        "100\n"
        "50\n"
        "01.15.26"
    )

async def generate_contract(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Генерация договора из данных"""
    text = update.message.text
    lines = text.split("\n")
    
    # Проверяем, что пользователь выбрал тип договора
    if 'contract_type' not in context.user_data:
        await update.message.reply_text(
            "⚠️ Сначала выбери тип договора через /start"
        )
        return
    
    # Проверяем, что ввел ровно 7 строк
    if len(lines) < 7:
        await update.message.reply_text(
            f"❌ Нужно ввести ровно 7 строк!\n"
            f"Ты ввел {len(lines)} строк.\n\n"
            "Отправь данные в таком порядке:\n"
            "1. Название бита\n"
            "2. Имя артиста\n"
            "3. Никнейм\n"
            "4. Адрес\n"
            "5. Цена (только цифры)\n"
            "6. Процент паблишинга (только цифры)\n"
            "7. Дата (в формате MM.DD.YY)"
        )
        return
    
    # Собираем данные в словарь
    data = {
        "beat": lines[0].strip(),
        "artist": lines[1].strip(),
        "nick": lines[2].strip(),
        "address": lines[3].strip(),
        "price": lines[4].strip(),
        "publishing": lines[5].strip(),
        "date": lines[6].strip()
    }
    
    # Проверяем, что цена и паблишинг - числа
    try:
        int(data["price"])
        int(data["publishing"])
    except ValueError:
        await update.message.reply_text(
            "❌ Цена и паблишинг должны быть числами!\n"
            "Попробуй еще раз."
        )
        return
    
    # Выбираем шаблон
    contract_type = context.user_data['contract_type']
    template_file = CONTRACT_TYPES[contract_type]["template"]
    
    # Генерируем договор
    publishing = int(data["publishing"])
    producer_publishing = 100 - publishing
    contract_date = data["date"] if data["date"] else datetime.now().strftime("%m.%d.%y")
    
    doc = DocxTemplate(template_file)
    doc.render({
        "BEAT": data["beat"],
        "ARTIST": data["artist"],
        "NICK": data["nick"],
        "ADDRESS": data["address"],
        "PRICE": data["price"],
        "PUBLISHING": publishing,
        "PRODUCER_PUBLISHING": producer_publishing,
        "DATE": contract_date
    })
    
    # Формируем название файла
    if contract_type == "exclusive":
        filename = f"Exclusive Rights for {data['beat']} beat for {data['artist']}.docx"
    else:
        filename = f"Unlimited Lease for {data['beat']} beat for {data['artist']}.docx"
    
    doc.save(filename)

    # 1. ВСЕГДА отправляем DOCX
    await update.message.reply_document(document=open(filename, "rb"))

    # 2. КОНВЕРТАЦИЯ В PDF через wkhtmltopdf
    try:
        pdf_filename = filename.replace('.docx', '.pdf')
        
        # Конвертируем DOCX в HTML
        html_content = docx_to_html(filename)
        
        # Указываем путь к wkhtmltopdf на Railway
        config = pdfkit.configuration(wkhtmltopdf='/usr/local/bin/wkhtmltopdf')
        
        # Создаём PDF из HTML
        pdfkit.from_string(html_content, pdf_filename, configuration=config)
        
        # Отправляем PDF
        await update.message.reply_document(document=open(pdf_filename, "rb"))
        os.remove(pdf_filename)
        print(f"PDF создан и отправлен: {pdf_filename}")
    except Exception as e:
        print(f"PDF не создан: {e}")
        await update.message.reply_text(
            "⚠️ Не удалось создать PDF, но DOCX отправлен."
        )

    # 3. Удаляем DOCX
    os.remove(filename)
    
    # Сбрасываем выбранный тип договора
    context.user_data['contract_type'] = None
    
    await update.message.reply_text(
        "✅ Готово! Договор отправлен.\n"
        "Чтобы создать новый - нажми /start"
    )

def main():
    app = Application.builder().token(TOKEN).build()
    
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CallbackQueryHandler(button_handler))
    app.add_handler(
        MessageHandler(
            filters.TEXT & ~filters.COMMAND,
            generate_contract
        )
    )
    
    print("Bot started...")
    app.run_polling()

if __name__ == "__main__":
    main()
